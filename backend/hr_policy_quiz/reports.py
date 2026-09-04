"""Generates the HR candidate report (summary + full question-by-question
answers for every completed candidate) as Excel, Word or PDF bytes."""
from io import BytesIO
from typing import List

from backend.hr_policy_quiz.db import candidates, attempts


def format_dt(dt):
    if not dt:
        return "-"
    return dt.strftime("%d %b %Y, %I:%M %p")


def gather_report_rows() -> List[dict]:
    """Same shape as /api/hr/candidates, plus the full answers_detail for
    completed candidates so the report can include a per-question review."""
    rows = []
    for c in candidates.find({}):
        latest_attempt = attempts.find_one(
            {"candidate_email": c["email"], "status": "completed"}, sort=[("submitted_at", -1)]
        )
        in_progress = attempts.find_one({"candidate_email": c["email"], "status": "in_progress"})
        rows.append(
            {
                "email": c["email"],
                "quiz_set_name": c.get("quiz_set_name") or "-",
                "status": "Completed" if latest_attempt else ("In progress" if in_progress else "Not started"),
                "score": latest_attempt["score"] if latest_attempt else None,
                "total_questions": latest_attempt["total_questions"] if latest_attempt else None,
                "submitted_at": latest_attempt["submitted_at"] if latest_attempt else None,
                "answers_detail": latest_attempt["answers_detail"] if latest_attempt else [],
            }
        )
    rows.sort(key=lambda r: r["email"])
    return rows


def result_label(item: dict) -> str:
    if item["is_correct"]:
        return "Correct"
    if item["selected_option"] is None:
        return "Unattempted"
    return "Incorrect"


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------

def generate_excel(rows: List[dict]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    bold = Font(bold=True)

    summary = wb.active
    summary.title = "Summary"
    summary.append(["Email", "Assigned quiz", "Status", "Score", "Total questions", "Submitted"])
    for cell in summary[1]:
        cell.font = bold
    for r in rows:
        summary.append(
            [r["email"], r["quiz_set_name"], r["status"], r["score"] or "", r["total_questions"] or "", format_dt(r["submitted_at"])]
        )

    answers = wb.create_sheet("Answers")
    answers.append(["Email", "Q#", "Question", "Candidate answer", "Correct answer", "Result"])
    for cell in answers[1]:
        cell.font = bold
    for r in rows:
        for i, item in enumerate(r["answers_detail"], start=1):
            selected = item["options"][item["selected_option"]] if item["selected_option"] is not None else "(no answer)"
            correct = item["options"][item["correct_index"]]
            answers.append([r["email"], i, item["question"], selected, correct, result_label(item)])

    for sheet in (summary, answers):
        for col in sheet.columns:
            width = max((len(str(c.value)) for c in col if c.value is not None), default=10)
            sheet.column_dimensions[col[0].column_letter].width = min(max(width + 2, 10), 60)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def generate_docx(rows: List[dict]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Candidate Report", level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Email", "Assigned quiz", "Status", "Score", "Total", "Submitted"]):
        cell.text = text
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["email"]
        cells[1].text = r["quiz_set_name"]
        cells[2].text = r["status"]
        cells[3].text = str(r["score"] or "-")
        cells[4].text = str(r["total_questions"] or "-")
        cells[5].text = format_dt(r["submitted_at"])

    for r in rows:
        if not r["answers_detail"]:
            continue
        doc.add_page_break()
        doc.add_heading(f"{r['email']} - {r['score']}/{r['total_questions']}", level=2)
        for i, item in enumerate(r["answers_detail"], start=1):
            selected = item["options"][item["selected_option"]] if item["selected_option"] is not None else "(no answer)"
            correct = item["options"][item["correct_index"]]
            p = doc.add_paragraph()
            p.add_run(f"Q{i}. {item['question']}").bold = True
            doc.add_paragraph(f"Candidate answer: {selected}")
            doc.add_paragraph(f"Correct answer: {correct}")
            doc.add_paragraph(f"Result: {result_label(item)}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def generate_pdf(rows: List[dict]) -> bytes:
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    def line(pdf, text, height=6):
        pdf.multi_cell(0, height, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "Candidate Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)

    col_widths = [45, 40, 25, 15, 20, 40]
    headers = ["Email", "Assigned quiz", "Status", "Score", "Total", "Submitted"]
    pdf.set_font("Helvetica", "B", 9)
    for w, h in zip(col_widths, headers):
        pdf.cell(w, 8, h, border=1)
    pdf.ln()
    pdf.set_font("Helvetica", "", 9)
    for r in rows:
        values = [r["email"], r["quiz_set_name"], r["status"], str(r["score"] or "-"), str(r["total_questions"] or "-"), format_dt(r["submitted_at"])]
        for w, v in zip(col_widths, values):
            pdf.cell(w, 8, v[:28], border=1)
        pdf.ln()

    for r in rows:
        if not r["answers_detail"]:
            continue
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 13)
        line(pdf, f"{r['email']} - {r['score']}/{r['total_questions']}", 10)
        pdf.set_font("Helvetica", "", 10)
        for i, item in enumerate(r["answers_detail"], start=1):
            selected = item["options"][item["selected_option"]] if item["selected_option"] is not None else "(no answer)"
            correct = item["options"][item["correct_index"]]
            pdf.set_font("Helvetica", "B", 10)
            line(pdf, f"Q{i}. {item['question']}")
            pdf.set_font("Helvetica", "", 10)
            line(pdf, f"Candidate answer: {selected}")
            line(pdf, f"Correct answer: {correct}")
            line(pdf, f"Result: {result_label(item)}")
            pdf.ln(2)

    return bytes(pdf.output())
